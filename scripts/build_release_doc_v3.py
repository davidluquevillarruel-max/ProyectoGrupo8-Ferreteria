from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "NOTA_VERSION_3.docx"


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    color = RGBColor(239, 119, 3) if level == 1 else RGBColor(52, 72, 86)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = color
        run.font.bold = True
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    table.rows[0].cells[0].text = "Elemento"
    table.rows[0].cells[1].text = "Detalle"
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    doc.add_paragraph()


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("Nota de version 3")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(18, 51, 22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    sub = subtitle.add_run("Funcionalidad global, data inicial y mini chat")
    sub.font.name = "Calibri"
    sub.font.size = Pt(13)
    sub.font.color.rgb = RGBColor(52, 72, 86)

    add_heading(doc, "Resumen", 1)
    doc.add_paragraph(
        "Esta version agrega acciones funcionales para botones principales, metricas administrativas "
        "iniciadas en cero, menu sticky al hacer scroll y un mini chat flotante para asesorar al usuario."
    )

    add_heading(doc, "Mejoras realizadas", 1)
    for item in [
        "Se agrego assets/js/ui-enhancements.js como script global.",
        "Los botones administrativos ahora exportan, guardan, agregan productos, crean usuarios demo o muestran accion contextual.",
        "Los enlaces sin destino real muestran una respuesta clara mediante toast.",
        "Las metricas admin inician en cero y se actualizan cuando se cargan productos o datos.",
        "La barra naranja de navegacion queda visible al hacer scroll con position: sticky.",
        "Se agrego un mini chat flotante con respuestas orientativas para asesorar al usuario.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Archivos modificados o creados", 1)
    add_table(doc, [
        ("assets/js/ui-enhancements.js", "Nuevo script global para botones, data, toast y mini chat."),
        ("assets/CSS/styles.css", "Estilos para menu sticky, toast, chat y estado vacio."),
        ("index.html", "Carga del script global y mini chat."),
        ("assets/pages/*.html", "Carga del script global en paginas internas y admin."),
        ("assets/pages/admin-*.html", "Metricas en cero y botones conectados a acciones."),
        ("docs/NOTA_VERSION_3.md", "Nota Markdown para commit."),
        ("docs/NOTA_VERSION_3.docx", "Word para adjuntar con la version."),
        ("scripts/build_release_doc_v3.py", "Script para regenerar este Word."),
    ])

    add_heading(doc, "Relacion con el silabo", 1)
    for item in [
        "HTML5/CSS3: mejora de interfaz, responsive y navegacion.",
        "JavaScript: comportamiento funcional en frontend.",
        "MVC, DAO, DTO y Facade: acciones preparadas para separarse en capas.",
        "JSP, Servlets y JDBC: la data local puede migrarse a backend Java.",
        "REST JSON: chat y panel admin quedan listos para consumir endpoints.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Verificacion", 1)
    for item in [
        "Mini chat visible y con respuesta en index.html.",
        "Menu naranja verificado con posicion sticky.",
        "Data administrativa verificada en cero antes de cargar productos.",
        "Producto de prueba agregado y contabilizado en inventario.",
        "Revision desktop y mobile sin desborde horizontal en paginas revisadas.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Mensaje sugerido para commit", 1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "feat: agregar funcionalidad global, data inicial y mini chat\n\n"
        "- Agregar script global ui-enhancements.js.\n"
        "- Hacer funcionales botones administrativos y enlaces sin accion.\n"
        "- Iniciar metricas admin en cero y actualizarlas al cargar productos.\n"
        "- Mantener visible el menu naranja con sticky scroll.\n"
        "- Agregar mini chat flotante para consultas con asesor.\n"
        "- Actualizar nota y Word para documentar la nueva version."
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
