from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "NOTA_VERSION_2.docx"

ORANGE = "EF7703"
DARK_GREEN = "123316"
BLUE_GRAY = "344856"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9DEE3"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + margin))
        if node is None:
            node = OxmlElement("w:" + margin)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(ORANGE if level == 1 else BLUE_GRAY)
        run.font.bold = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    return p


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in (
        ("Title", 22, DARK_GREEN),
        ("Heading 1", 16, ORANGE),
        ("Heading 2", 13, BLUE_GRAY),
        ("Heading 3", 12, BLUE_GRAY),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_key_value_table(doc):
    table = doc.add_table(rows=5, cols=2)
    set_cell_margins(table)
    set_table_width(table, [2600, 6760])
    rows = [
        ("Proyecto", "Estructuras & Diseños Group - Ferretería"),
        ("Versión", "Versión 2"),
        ("Curso", "Desarrollo Web Integrado"),
        ("Enfoque", "Frontend HTML5/CSS3/JavaScript preparado para Java EE"),
        ("Fecha", "23 de mayo de 2026"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell in row.cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()


def add_files_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_cell_margins(table)
    set_table_width(table, [2600, 2000, 4760])
    headers = ("Archivo", "Tipo", "Detalle")
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_border(cell)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)

    rows = [
        ("index.html", "Modificado", "Conexión de accesos del panel admin desde la hamburguesa y corrección de enlace roto."),
        ("assets/CSS/styles.css", "Modificado", "Estilos para Sobre Nosotros, panel admin y responsive móvil."),
        ("assets/pages/nosotros.html", "Rehecho", "Contenido institucional y estructura visual de ferretería."),
        ("assets/pages/admin-inventario.html", "Nuevo", "Control visual de stock, métricas y tabla de inventario."),
        ("assets/pages/admin-ventas.html", "Nuevo", "Indicadores de ventas, ranking y ventas por categoría."),
        ("assets/pages/admin-usuarios.html", "Nuevo", "Control visual de usuarios, roles y estados."),
        ("assets/pages/admin-productos.html", "Nuevo", "Formulario visual para agregar, ocultar o quitar productos."),
        ("README.md", "Nuevo", "Estructura del proyecto, páginas principales y ubicación de documentación."),
        ("docs/NOTA_VERSION_2.md", "Nuevo", "Resumen en Markdown para acompañar el commit."),
        ("docs/NOTA_VERSION_2.docx", "Nuevo", "Documento Word formal con resumen de versión para adjuntar."),
        ("docs/qa/", "Nuevo", "Carpeta con capturas de verificación desktop y móvil."),
        ("scripts/build_release_doc.py", "Nuevo", "Script que genera el documento Word de la nota de versión."),
    ]
    for file_name, kind, detail in rows:
        cells = table.add_row().cells
        values = (file_name, kind, detail)
        for cell, value in zip(cells, values):
            cell.text = value
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Nota de versión 2")
    run.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Proyecto Ferretería - Estructuras & Diseños Group")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor.from_string(BLUE_GRAY)
    subtitle.paragraph_format.space_after = Pt(18)

    add_key_value_table(doc)

    add_heading(doc, "Resumen ejecutivo", 1)
    add_body(
        doc,
        "Esta versión mejora la base visual y funcional del proyecto web de ferretería, "
        "manteniendo el alcance del sílabo de Desarrollo Web Integrado. Las mejoras se "
        "concentran en una página institucional, un panel de gestión para administrador, "
        "enlaces funcionales desde la hamburguesa y una capa responsive para celulares."
    )

    add_heading(doc, "Mejoras realizadas", 1)
    add_heading(doc, "Página Sobre Nosotros", 2)
    add_bullet(doc, "Se reemplazó la página anterior que tenía contenido de otra temática y textos de relleno.")
    add_bullet(doc, "Se incorporó contenido institucional alineado a una ferretería.")
    add_bullet(doc, "Se agregaron secciones de propósito, valores, métricas y llamados a la acción.")

    add_heading(doc, "Panel de gestión para administrador", 2)
    add_bullet(doc, "Se crearon páginas para inventario, análisis de ventas, control de usuarios y gestión de productos.")
    add_bullet(doc, "Se diseñaron tarjetas de métricas, tablas, formularios, estados y acciones rápidas.")
    add_bullet(doc, "Se enlazaron las opciones del panel admin desde la hamburguesa del inicio.")

    add_heading(doc, "Responsive móvil", 2)
    add_bullet(doc, "Se ajustaron grillas, botones, encabezados, tarjetas, sidebar admin y tablas.")
    add_bullet(doc, "Se validó que no exista desborde horizontal en las vistas móviles revisadas.")
    add_bullet(doc, "Las tablas administrativas mantienen legibilidad con desplazamiento interno cuando el contenido es ancho.")

    add_heading(doc, "Orden del proyecto e imágenes", 2)
    add_bullet(doc, "Se creó README.md para documentar la estructura del proyecto y páginas principales.")
    add_bullet(doc, "Se centralizó la documentación y las capturas QA dentro de docs/.")
    add_bullet(doc, "Se incorporaron imágenes reales desde assets/img herramientas/ en inventario, ranking de ventas y gestión de productos.")

    add_heading(doc, "Relación con el sílabo", 1)
    add_body(
        doc,
        "Los cambios corresponden a la base frontend del proyecto y se mantienen dentro del marco del curso. "
        "La implementación actual refuerza HTML5, CSS3 y JavaScript, y deja preparada la evolución hacia "
        "JSP, Servlets, JDBC, MVC, DAO, DTO, Façade, JSF y RESTful APIs con JSON."
    )
    add_bullet(doc, "Unidad 1: estructura web con HTML5, CSS3 y preparación para arquitectura Java EE.")
    add_bullet(doc, "Unidad 2: separación futura de vistas, controladores y acceso a datos mediante patrones.")
    add_bullet(doc, "Unidad 3: base visual lista para integrarse con JSP, JSF o frameworks frontend.")
    add_bullet(doc, "Unidad 4: panel admin preparado conceptualmente para APIs REST con JSON.")

    add_heading(doc, "Archivos modificados o creados", 1)
    add_files_table(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Verificación realizada", 1)
    add_bullet(doc, "Servidor local ejecutado en http://127.0.0.1:5501.")
    add_bullet(doc, "Carga verificada de la página Sobre Nosotros.")
    add_bullet(doc, "Carga verificada de las cuatro páginas administrativas.")
    add_bullet(doc, "Navegación verificada desde la hamburguesa hacia Inventario.")
    add_bullet(doc, "Revisión responsive móvil en inicio, productos, Sobre Nosotros, Inventario y Productos admin.")
    add_bullet(doc, "Verificación de carga de imágenes reales en Inventario, Analizar venta y Agregar o quitar productos.")
    add_bullet(doc, "Resultado: sin desborde horizontal en las vistas revisadas.")

    add_heading(doc, "Mensaje sugerido para commit", 1)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        "feat: mejorar sobre nosotros, panel admin y responsive\n\n"
        "- Rehacer la página Sobre Nosotros con contenido de ferretería.\n"
        "- Crear páginas admin para inventario, ventas, usuarios y productos.\n"
        "- Conectar accesos del panel de gestión desde la hamburguesa.\n"
        "- Agregar estilos desktop y responsive para celular.\n"
        "- Ordenar documentación en docs/ y agregar README.\n"
        "- Usar imágenes reales del proyecto en páginas admin.\n"
        "- Corregir enlace roto hacia servicios.html.\n"
        "- Mantener la estructura preparada para integración Java EE según el sílabo."
    )
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(BLUE_GRAY)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
